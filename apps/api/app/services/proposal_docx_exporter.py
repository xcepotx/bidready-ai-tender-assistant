from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from app.services.i18n_service import localized_label


BRAND_NAVY = RGBColor(11, 29, 58)
BRAND_TEAL = RGBColor(0, 168, 142)
BRAND_SLATE = RGBColor(71, 85, 105)


def _set_run_style(run, *, bold=False, size=None, color=None):
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_style(run, bold=True, size=24, color=BRAND_NAVY)


def _add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_style(run, size=10, color=BRAND_SLATE)


def _add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        _set_run_style(run, bold=True, size=16, color=BRAND_NAVY)
    else:
        run = p.add_run(text)
        _set_run_style(run, bold=True, size=12, color=BRAND_NAVY)


def _add_bullet(doc, text):
    if not text:
        return
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(str(text))


def _add_numbered(doc, text):
    if not text:
        return
    p = doc.add_paragraph(style="List Number")
    p.add_run(str(text))


def _add_field_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value or "-")

        for run in cells[0].paragraphs[0].runs:
            _set_run_style(run, bold=True, size=9, color=BRAND_NAVY)

        for run in cells[1].paragraphs[0].runs:
            _set_run_style(run, size=9, color=BRAND_SLATE)

    doc.add_paragraph()




def _template_value(template, field, default=None):
    if not template:
        return default
    value = getattr(template, field, None)
    return value if value not in (None, "") else default


def _ordered_sections(proposal_sections, template):
    sections = list(proposal_sections or [])
    excluded = set(_template_value(template, "excluded_section_keys", []) or [])
    order = _template_value(template, "section_order", []) or []

    if excluded:
        sections = [section for section in sections if section.section_key not in excluded]

    if order:
        order_index = {key: idx for idx, key in enumerate(order)}
        sections.sort(key=lambda section: (order_index.get(section.section_key, 999), section.section_order, section.id))

    return sections


def _add_template_block(doc, title, content):
    if not content:
        return
    _add_heading(doc, title, level=2)
    for paragraph in str(content).splitlines():
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())


def _add_custom_sections(doc, custom_sections):
    for index, section in enumerate(custom_sections or [], start=1):
        if not isinstance(section, dict):
            continue

        title = section.get("title") or section.get("section_title") or f"Custom Section {index}"
        content = section.get("content") or section.get("draft_content") or section.get("body") or ""
        if not title and not content:
            continue

        _add_heading(doc, str(title), level=1)
        for paragraph in str(content).splitlines():
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        doc.add_page_break()

def _join_list(items):
    if not items:
        return "-"
    return "\n".join([f"• {item}" for item in items if item])


def export_proposal_draft_docx(
    *,
    project,
    metadata,
    proposal_sections,
    response_items,
    output_path: str,
    evidence_items=None,
    decision_gate=None,
    output_language: str = "en",
    proposal_template=None,
) -> str:
    evidence_items = evidence_items or []
    L = lambda label: localized_label(label, output_language)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    template_name = _template_value(proposal_template, "template_name", "Standard Executive Proposal")
    executive_title = _template_value(proposal_template, "executive_title", "BidReady AI - Proposal Draft")
    cover_note = _template_value(proposal_template, "cover_note")
    company_profile = _template_value(proposal_template, "company_profile")
    win_theme = _template_value(proposal_template, "win_theme")
    proposal_tone = _template_value(proposal_template, "proposal_tone", "formal")
    custom_sections = _template_value(proposal_template, "custom_sections", []) or []
    footer_note = _template_value(proposal_template, "footer_note")

    package_name = metadata.package_name if metadata else project.title
    issuer = metadata.issuer if metadata and metadata.issuer else project.issuer
    service_domain = metadata.service_domain if metadata and metadata.service_domain else project.tender_type

    _add_title(doc, executive_title)
    _add_subtitle(doc, "Tender Intelligence Platform")
    _add_subtitle(doc, "Generated from proposal outline, response plan, evidence checklist, and RFP metadata.")
    doc.add_paragraph()
    _add_template_block(doc, "Cover Note", cover_note)
    _add_template_block(doc, "Company Profile", company_profile)
    _add_template_block(doc, "Win Theme", win_theme)
    if cover_note or company_profile or win_theme:
        doc.add_paragraph()

    _add_heading(doc, L("Tender Metadata"), level=1)
    _add_field_table(doc, [
        ("Project", project.title),
        ("Package Name", package_name),
        ("Issuer", issuer),
        ("Service Domain", service_domain),
        ("Submission Deadline", metadata.submission_deadline if metadata else None),
        ("Clarification Deadline", metadata.clarification_deadline if metadata else None),
        ("Proposal Validity", metadata.proposal_validity if metadata else None),
        ("Generation Mode", "rules_only"),
    ])

    if metadata and metadata.submission_requirements:
        _add_heading(doc, L("Submission Requirements"), level=2)
        for item in metadata.submission_requirements:
            _add_bullet(doc, item)
        doc.add_paragraph()

    _add_heading(doc, L("Executive Decision Gate"), level=1)

    if decision_gate:
        _add_field_table(doc, [
            (L("Decision Status"), decision_gate.decision_status),
            (L("Recommendation"), decision_gate.recommendation),
            (L("Readiness Score"), decision_gate.readiness_score),
            (L("Confidence"), decision_gate.confidence),
            (L("Owner"), decision_gate.owner),
            (L("Due Date"), decision_gate.due_date),
            (L("Generation Mode"), decision_gate.generation_mode),
        ])

        if decision_gate.executive_summary:
            _add_heading(doc, L("Executive Summary"), level=2)
            doc.add_paragraph(decision_gate.executive_summary)

        decision_sections = [
            (L("Key Reasons"), decision_gate.key_reasons or []),
            (L("Blockers"), decision_gate.blockers or []),
            (L("Required Approvals"), decision_gate.required_approvals or []),
            (L("Next Actions"), decision_gate.next_actions or []),
        ]

        for title, items in decision_sections:
            _add_heading(doc, title, level=2)
            if items:
                for item in items:
                    _add_bullet(doc, item)
            else:
                doc.add_paragraph("-")

        if decision_gate.notes:
            _add_heading(doc, L("Decision Notes"), level=2)
            doc.add_paragraph(decision_gate.notes)
    else:
        doc.add_paragraph(L("Decision gate has not been generated yet."))

    doc.add_page_break()

    _add_heading(doc, L("Proposal Sections"), level=1)

    ordered_sections = _ordered_sections(proposal_sections, proposal_template)

    for section_item in ordered_sections:
        _add_heading(doc, f"{section_item.section_order}. {section_item.title}", level=1)

        _add_field_table(doc, [
            ("Owner", section_item.owner),
            ("Status", section_item.status),
            ("Section Key", section_item.section_key),
            ("Source Response Item IDs", ", ".join([str(x) for x in section_item.source_response_item_ids or []])),
        ])

        if section_item.purpose:
            _add_heading(doc, "Purpose", level=2)
            doc.add_paragraph(section_item.purpose)

        if section_item.content_outline:
            _add_heading(doc, "Content Outline", level=2)
            for item in section_item.content_outline:
                _add_bullet(doc, item)

        if section_item.draft_content:
            _add_heading(doc, "Draft Content", level=2)
            for paragraph in str(section_item.draft_content).splitlines():
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

        if section_item.evidence_needed:
            _add_heading(doc, "Evidence Needed", level=2)
            for item in section_item.evidence_needed:
                _add_bullet(doc, item)

        if section_item.risks:
            _add_heading(doc, "Risks / Review Notes", level=2)
            for item in section_item.risks:
                _add_bullet(doc, item)

        if section_item.notes:
            _add_heading(doc, "Section Notes", level=2)
            doc.add_paragraph(section_item.notes)

        doc.add_page_break()

    # Proposal template custom sections.
    _add_custom_sections(doc, custom_sections)

    _add_heading(doc, L("Response Plan Appendix"), level=1)

    if not response_items:
        doc.add_paragraph("No response plan items available.")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        headers = [
            "ID",
            "Category",
            "Compliance",
            "Owner",
            "Status",
            "Requirement",
        ]

        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        for item in response_items:
            cells = table.add_row().cells
            cells[0].text = str(item.id)
            cells[1].text = item.category or "-"
            cells[2].text = item.compliance_status or "-"
            cells[3].text = item.owner or "-"
            cells[4].text = item.status or "-"
            cells[5].text = item.requirement_text or "-"

    doc.add_page_break()

    _add_heading(doc, L("Evidence Checklist Appendix"), level=1)

    if evidence_items:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Evidence"
        table.rows[0].cells[1].text = "Category"
        table.rows[0].cells[2].text = "Owner"
        table.rows[0].cells[3].text = "Priority"
        table.rows[0].cells[4].text = "Status"

        for evidence in evidence_items:
            cells = table.add_row().cells
            cells[0].text = evidence.evidence_name or "-"
            cells[1].text = evidence.evidence_category or "-"
            cells[2].text = evidence.owner or "-"
            cells[3].text = evidence.priority or "-"
            cells[4].text = evidence.status or "-"
    else:
        evidence_map = {}

        for item in response_items:
            for evidence in item.evidence_needed or []:
                evidence_map.setdefault(evidence, set()).add(item.owner or "unassigned")

        if not evidence_map:
            doc.add_paragraph("No evidence checklist available.")
        else:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Evidence"
            table.rows[0].cells[1].text = "Suggested Owner(s)"
            table.rows[0].cells[2].text = "Status"

            for evidence, owners in sorted(evidence_map.items()):
                cells = table.add_row().cells
                cells[0].text = evidence
                cells[1].text = ", ".join(sorted(owners))
                cells[2].text = "open"

    if footer_note:
        _add_heading(doc, "Template Footer Note", level=2)
        doc.add_paragraph(str(footer_note))

    doc.save(output)
    return str(output)
